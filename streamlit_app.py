import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import html as html_lib
import json as _json
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import os

st.set_page_config(
    page_title="Gene Coverage Analyzer",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ── Stitch Design System ───────────────────────────────────────────────────
# CSS is injected via JS (components.html) so Streamlit Cloud cannot strip
# the <style> tag the way it does with st.markdown.
_GCA_CSS = """
/* ── Tokens ──────────────────────────────────────────────────────────────── */
:root{
  --cp:        #4f46e5;
  --cp-dk:     #4338ca;
  --cs:        #64748b;
  --bg:        #f5f7fb;
  --surf:      #ffffff;
  --surf-lo:   #f1f5f9;
  --surf-mid:  #e9eef6;
  --surf-hi:   #eef2ff;
  --surf-max:  #e0e7ff;
  --out:       #cbd5e1;
  --out-v:     #e2e8f0;
  --on-surf:   #0f172a;
  --err:       #dc2626;
  --err-c:     #fee2e2;
  --err-on:    #991b1b;
  --r-card:    14px;
  --r-pill:    100px;
  --sh:        0 1px 3px rgba(15,23,42,.06);
  --sh-hover:  0 8px 24px rgba(15,23,42,.10);
}

/* ── Base ────────────────────────────────────────────────────────────────── */
html,body,[class*="css"],.stMarkdown{font-family:'Inter',sans-serif!important}
.stApp{background:var(--bg)!important}
.block-container{
  padding-top:0!important;
  padding-left:2rem!important;padding-right:2rem!important;
  max-width:1280px!important;
}

/* ── Hide chrome ─────────────────────────────────────────────────────────── */
#MainMenu,footer,[data-testid="stToolbar"],[data-testid="stDecoration"],.stDeployButton{display:none!important}
header[data-testid="stHeader"]{display:none!important}

/* ── Top navbar ──────────────────────────────────────────────────────────── */
.gca-nav{
  background:var(--surf);border-bottom:1px solid var(--out-v);
  height:60px;padding:0 1.5rem;
  display:flex;align-items:center;justify-content:space-between;
  margin:0 -2rem 1.75rem -2rem;
  box-shadow:0 1px 3px rgba(0,0,0,.06);
  position:sticky;top:0;z-index:999;
}
.gca-brand{display:flex;align-items:center;gap:10px}
.gca-brand-icon{
  font-family:'Material Symbols Outlined';font-size:26px;color:var(--cp);
  font-variation-settings:'FILL' 0,'wght' 400,'GRAD' 0,'opsz' 24;
}
.gca-brand-name{font-size:17px;font-weight:700;color:var(--cp);line-height:1.15;display:block}
.gca-brand-sub{font-size:11px;color:var(--cs);letter-spacing:.02em;display:block}
.gca-nav-links{display:flex;gap:1.25rem}
.gca-nav-links a{
  font-size:13px;color:var(--cs);text-decoration:none;
  padding-bottom:2px;border-bottom:2px solid transparent;transition:color .15s;
}
.gca-nav-links a:hover{color:var(--cp)}
.gca-nav-links a.active{color:var(--cp);border-bottom-color:var(--cp);font-weight:600}

/* ── Section label ───────────────────────────────────────────────────────── */
.gca-lbl{
  text-align:center;font-size:15px;font-weight:600;color:var(--on-surf);
  margin-bottom:.75rem;display:block;width:100%;
}

/* ── Step-1 centering wrapper ────────────────────────────────────────────── */
.gca-step1{
  display:flex;flex-direction:column;align-items:center;
  padding:1.5rem 0 1.25rem;
}

/* ── Radio → pill toggle ─────────────────────────────────────────────────── */
div[data-testid="stRadio"]{
  text-align:center!important;
  width:100%!important;
  margin-bottom:.5rem!important;
}
div[data-testid="stRadio"]>div{
  background:var(--surf-hi);border-radius:var(--r-pill);
  border:1px solid var(--out-v);padding:3px;
  gap:0!important;display:inline-flex!important;
}
div[data-testid="stRadio"]>div>label{
  border-radius:var(--r-pill);padding:4px 20px!important;
  margin:0!important;font-size:13px;font-weight:400;color:var(--cs);
  transition:all .15s;cursor:pointer;
}
div[data-testid="stRadio"]>div>label:has(input:checked){
  background:var(--surf)!important;
  box-shadow:0 1px 3px rgba(0,0,0,.14);
  color:var(--cp)!important;font-weight:600;
}
div[data-testid="stRadio"]>div>label>div:first-child{display:none!important}

/* ── Upload column subtle separator ─────────────────────────────────────── */
/* Columns are separated by gap only — no card border (avoids empty-box artifact) */

/* ── Buttons ─────────────────────────────────────────────────────────────── */
.stButton>button{
  background:linear-gradient(135deg,var(--cp),var(--cp-dk))!important;color:#fff!important;
  border:none!important;border-radius:10px!important;
  font-family:'Inter',sans-serif!important;font-size:14px!important;font-weight:600!important;
  padding:.55rem 1.25rem!important;letter-spacing:.01em;
  box-shadow:0 2px 6px rgba(79,70,229,.28)!important;
  transition:transform .15s,box-shadow .2s,filter .2s!important;
}
.stButton>button:hover{filter:brightness(1.06)!important;box-shadow:0 6px 16px rgba(79,70,229,.34)!important;transform:translateY(-1px)!important}
.stButton>button:active{transform:translateY(0)!important;filter:brightness(.96)!important}
.gca-sec .stButton>button{
  background:var(--surf)!important;color:var(--on-surf)!important;
  border:1px solid var(--out)!important;box-shadow:none!important;
}
.gca-sec .stButton>button:hover{border-color:var(--cp)!important;color:var(--cp)!important;opacity:1!important}

/* ── Download buttons ────────────────────────────────────────────────────── */
.stDownloadButton>button{
  background:var(--surf)!important;color:var(--on-surf)!important;
  border:1px solid var(--out)!important;border-radius:8px!important;
  font-family:'Inter',sans-serif!important;font-size:14px!important;font-weight:500!important;
  box-shadow:none!important;
}
.stDownloadButton>button:hover{border-color:var(--cp)!important;color:var(--cp)!important}

/* ── File uploader ───────────────────────────────────────────────────────── */
[data-testid="stFileUploader"] section{
  border:2px dashed var(--out)!important;border-radius:10px!important;
  background:white!important;transition:border-color .2s,background .2s!important;
}
[data-testid="stFileUploader"] section:hover{
  border-color:var(--cp)!important;background:rgba(79,70,229,.03)!important;
}

/* ── Tabs ────────────────────────────────────────────────────────────────── */
.stTabs [data-baseweb="tab-list"]{
  border-bottom:2px solid var(--out-v);gap:4px;background:transparent!important;
}
.stTabs [data-baseweb="tab"]{
  font-family:'Inter',sans-serif!important;font-size:13px!important;
  color:var(--cs)!important;padding:6px 12px!important;
  background:transparent!important;border-bottom:2px solid transparent;
}
.stTabs [data-baseweb="tab"][aria-selected="true"]{
  color:var(--cp)!important;border-bottom-color:var(--cp)!important;
  font-weight:600!important;background:transparent!important;
}
.stTabs [data-baseweb="tab-highlight"],.stTabs [data-baseweb="tab-border"]{display:none!important}

/* ── Expanders ───────────────────────────────────────────────────────────── */
[data-testid="stExpander"]{
  border:1px solid var(--out-v)!important;border-radius:var(--r-card)!important;
  background:var(--surf)!important;box-shadow:var(--sh)!important;overflow:hidden;
}
[data-testid="stExpander"] summary{
  font-family:'Inter',sans-serif!important;font-size:15px!important;
  font-weight:600!important;color:var(--on-surf)!important;
  background:var(--surf)!important;padding:1rem 1.25rem!important;
}
[data-testid="stExpander"] summary:hover{background:var(--surf-lo)!important}
[data-testid="stExpander"] summary svg{color:var(--cs)!important}

/* ── Text area ───────────────────────────────────────────────────────────── */
.stTextArea>div>div>textarea{
  font-family:'JetBrains Mono',monospace!important;font-size:13px!important;
  border-radius:8px!important;border-color:var(--out)!important;background:white!important;
}
.stTextArea>div>div>textarea:focus{
  border-color:var(--cp)!important;
  box-shadow:0 0 0 3px rgba(79,70,229,.14)!important;outline:none!important;
}

/* ── Headings ────────────────────────────────────────────────────────────── */
h3{font-size:15px!important;font-weight:600!important;color:var(--on-surf)!important}

/* ── Custom metric cards ─────────────────────────────────────────────────── */
.gca-metrics{display:grid;grid-template-columns:repeat(3,1fr);gap:1rem;margin:1.25rem 0}
.gca-metric{
  background:white;border:1px solid var(--out-v);border-radius:14px;
  padding:1.3rem 1.5rem;display:flex;flex-direction:column;
  box-shadow:var(--sh);transition:transform .15s,box-shadow .2s;
  position:relative;overflow:hidden;
}
.gca-metric::before{
  content:'';position:absolute;top:0;left:0;width:100%;height:3px;
  background:linear-gradient(90deg,var(--cp),var(--cp-dk));
}
.gca-metric:hover{transform:translateY(-2px);box-shadow:var(--sh-hover)}
.gca-metric.err{
  background:linear-gradient(180deg,rgba(220,38,38,.04),#fff);border-color:var(--err-c);
}
.gca-metric.err::before{background:linear-gradient(90deg,var(--err),#f87171)}
.gca-metric-lbl{
  font-size:10px;font-weight:700;letter-spacing:.06em;text-transform:uppercase;
  color:var(--cs);margin-bottom:.4rem;
}
.gca-metric.err .gca-metric-lbl{color:var(--err-on)}
.gca-metric-val{font-size:34px;font-weight:700;letter-spacing:-.02em;line-height:1.1;color:var(--on-surf)}
.gca-metric.err .gca-metric-val{color:var(--err)}

/* ── Results header ──────────────────────────────────────────────────────── */
.gca-rh{
  display:flex;justify-content:space-between;align-items:flex-end;
  padding-bottom:1rem;border-bottom:1px solid var(--out-v);margin-bottom:1.25rem;
}
.gca-rh-ttl{font-size:22px;font-weight:600;color:var(--on-surf);line-height:1.2}
.gca-rh-sub{font-size:12px;color:var(--cs);margin-top:2px}

/* ── Gene badges ─────────────────────────────────────────────────────────── */
.gca-gene-grid{display:flex;flex-wrap:wrap;gap:6px;max-height:320px;overflow-y:auto}
.gca-gene{
  display:inline-flex;align-items:center;padding:3px 12px;
  border-radius:var(--r-pill);font-family:'JetBrains Mono',monospace;font-size:12px;
  transition:transform .1s;cursor:default;
}
.gca-gene:hover{transform:translateY(-1px)}
.gca-gene.ok{background:var(--surf-hi);color:var(--cp-dk);border:1px solid var(--surf-max)}
.gca-gene.low{
  background:var(--err-c);color:var(--err-on);
  border:1px solid rgba(220,38,38,.30);font-weight:600;
}
.gca-gene .pct{font-size:10px;opacity:.7;margin-left:4px}

/* ── Card section wrapper ────────────────────────────────────────────────── */
.gca-card{
  background:var(--surf);border:1px solid var(--out-v);
  border-radius:var(--r-card);box-shadow:var(--sh);padding:1.5rem;
}
.gca-card-title{
  font-size:16px;font-weight:600;color:var(--on-surf);
  display:flex;align-items:center;gap:.6rem;margin-bottom:1rem;min-height:26px;
}
.gca-badge-ready{
  font-size:11px;font-weight:600;background:rgba(5,150,105,.10);color:#047857;
  border:1px solid rgba(5,150,105,.22);padding:2px 10px;border-radius:100px;
  display:inline-flex;align-items:center;gap:6px;
}
.gca-badge-ready::before{content:'';width:6px;height:6px;border-radius:50%;background:#059669}

/* ── Status / progress / divider ─────────────────────────────────────────── */
.stSuccess>div,.stError>div,.stInfo>div,.stWarning>div{border-radius:8px!important}
/* progress: light track + indigo gradient fill (single clean bar) */
[data-testid="stProgress"] div[role="progressbar"]>div{
  background:var(--surf-mid)!important;border-radius:100px!important;height:8px!important;
}
[data-testid="stProgress"] div[role="progressbar"]>div>div{
  background:linear-gradient(90deg,var(--cp),var(--cp-dk))!important;
  border-radius:100px!important;
}
[data-testid="stProgress"] p{color:var(--cs)!important;font-size:13px!important}
.stSpinner>div{border-top-color:var(--cp)!important}
hr{border-color:var(--out-v)!important;margin:1.5rem 0!important}

/* ── Footer bar ──────────────────────────────────────────────────────────── */
.gca-footer{
  text-align:center;padding:.75rem 1rem;background:var(--surf-max);
  border-top:1px solid var(--out-v);font-size:11px;color:var(--cs);
  display:flex;align-items:center;justify-content:center;gap:6px;
  margin:2rem -2rem -2rem -2rem;
}
.gca-footer .ms{
  font-family:'Material Symbols Outlined';font-size:14px;
  font-variation-settings:'FILL' 0,'wght' 400,'GRAD' 0,'opsz' 24;
}

@media(max-width:768px){
  .gca-metrics{grid-template-columns:1fr}
  .gca-nav-links{display:none}
  .block-container{padding-left:1rem!important;padding-right:1rem!important}
}
"""

_FONT_URLS = [
    "https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700"
    "&family=JetBrains+Mono:wght@400&display=swap",
    "https://fonts.googleapis.com/css2?family=Material+Symbols+Outlined"
    ":opsz,wght,FILL,GRAD@24,400,0,0",
]

components.html(
    f"""<script>
(function(){{
  var p = window.parent.document;
  if (p.getElementById('gca-css')) return;
  {_json.dumps(_FONT_URLS)}.forEach(function(href){{
    var l = p.createElement('link');
    l.rel='stylesheet'; l.href=href;
    p.head.appendChild(l);
  }});
  var s = p.createElement('style');
  s.id='gca-css';
  s.textContent={_json.dumps(_GCA_CSS)};
  p.head.appendChild(s);
}})();
</script>""",
    height=0,
    scrolling=False,
)

# ── Top Nav Bar ────────────────────────────────────────────────────────────
st.markdown("""
<header class="gca-nav">
  <div class="gca-brand">
    <span class="gca-brand-icon">science</span>
    <div>
      <span class="gca-brand-name">Gene Coverage Analyzer</span>
      <span class="gca-brand-sub">Clinical Genomics Dashboard</span>
    </div>
  </div>
</header>
""", unsafe_allow_html=True)

# ── Session State ──────────────────────────────────────────────────────────
if 'coverage_data' not in st.session_state:
    st.session_state.coverage_data = None
if 'panel_genes' not in st.session_state:
    st.session_state.panel_genes = []
if 'filtered_data' not in st.session_state:
    st.session_state.filtered_data = None
if 'file_basename' not in st.session_state:
    st.session_state.file_basename = "output"
if 'last_processed_file' not in st.session_state:
    st.session_state.last_processed_file = None
if 'raw_batch' not in st.session_state:
    st.session_state.raw_batch = {}
if 'raw_batch_sig' not in st.session_state:
    st.session_state.raw_batch_sig = None
if 'mito_data' not in st.session_state:
    st.session_state.mito_data = None
if 'pipeline_html_bytes' not in st.session_state:
    st.session_state.pipeline_html_bytes = None
if 'pipeline_html_name' not in st.session_state:
    st.session_state.pipeline_html_name = None

# ── Helper functions ───────────────────────────────────────────────────────
@st.cache_data
def extract_gene_name(name):
    if isinstance(name, str):
        if ',' in name:
            return name.split(',')[0]
        elif ';' in name:
            return name.split(';')[0]
        return name
    return None

def parse_gene_list(text):
    """Returns (unique_genes, n_duplicates)."""
    if not text.strip():
        return [], 0
    text = text.replace('\n', ' ').replace(',', ' ').replace('\t', ' ')
    genes = [g.strip() for g in text.split() if g.strip()]
    seen, unique = set(), []
    for g in genes:
        if g not in seen:
            seen.add(g)
            unique.append(g)
    return unique, len(genes) - len(unique)

@st.cache_data(show_spinner=False)
def preprocess_excel_cached(file_bytes, file_name, skip_rows=1):
    try:
        data = pd.read_excel(io.BytesIO(file_bytes), skiprows=skip_rows)
        if 'Gene Name' not in data.columns:
            raise ValueError("Column 'Gene Name' not found")
        data['Gene_Name'] = data['Gene Name'].apply(extract_gene_name)
        data['Gene_ID'] = data['Gene_Name'].str.strip().str.upper()
        valid = (
            data['Gene_ID'].notna() & (data['Gene_ID'] != '') &
            ~data['Gene_ID'].str.startswith('ENSG') &
            ~data['Gene_ID'].str.startswith('INTRON')
        )
        data = data[valid].copy()
        data['Covered_Bases'] = (data['% 1x'] / 100.0) * data['Counted Bases']
        grouped = data.groupby('Gene_ID', as_index=False).agg(
            Total_Bases=('Counted Bases', 'sum'),
            Covered_Bases=('Covered_Bases', 'sum'),
        )
        grouped['% 1x'] = (grouped['Covered_Bases'] / grouped['Total_Bases'] * 100).round(2)
        grouped['Gene_Name'] = grouped['Gene_ID']
        grouped['Ref Name'] = grouped['Gene_ID']
        return grouped, os.path.splitext(file_name)[0]
    except Exception as e:
        raise Exception(f"Error: {str(e)}")

def create_word_document_with_mito(df_data):
    df_data = df_data.reset_index(drop=True)
    doc = Document()
    doc.add_heading('Appendix 1: Gene Coverage', 1)
    doc.add_heading('Indication Based Analysis:', 2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)

    chunk_size = 4
    chunks = [df_data[i:i+chunk_size] for i in range(0, len(df_data), chunk_size)]
    table = doc.add_table(rows=1, cols=chunk_size * 2)
    table.alignment = 1
    hdr_cells = table.rows[0].cells

    for i in range(chunk_size):
        for cell, text in [
            (hdr_cells[i*2],     "Gene Name"),
            (hdr_cells[i*2+1],   "Percentage of coding region covered"),
        ]:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            run.font.name = 'Calibri'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            run.font.size = Pt(8)

    for chunk in chunks:
        chunk = chunk.reset_index(drop=True)
        row_cells = table.add_row().cells
        for i in range(chunk_size):
            if i < len(chunk):
                row = chunk.iloc[i]
                gene = str(row['Gene_ID'])
                pv = row['% 1x']
                percent = float(pv.iloc[0] if isinstance(pv, pd.Series) else pv) if pd.notna(pv) else 0.0

                for cell, content, italic in [
                    (row_cells[i*2],   gene,         True),
                    (row_cells[i*2+1], str(percent), False),
                ]:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    r = p.add_run(content)
                    r.font.name = 'Calibri'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    r.font.size = Pt(8)
                    r.italic = italic
                    if percent < 90:
                        r.font.color.rgb = RGBColor(255, 0, 0)
            else:
                for cell in [row_cells[i*2], row_cells[i*2+1]]:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    r = p.add_run("–")
                    r.font.name = 'Calibri'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(255, 255, 255)

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for bn in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    return doc

def generate_html_report(df, patient_name):
    gene_col = next(
        (c for c in ['Gene_ID', 'Gene_Name', 'Gene', 'gene'] if c in df.columns),
        df.columns[0]
    )
    coverage_col = next(
        (c for c in ['% 1x', '% 1x ', 'Perc_1x', 'Coverage', 'Percentage'] if c in df.columns),
        df.columns[1]
    )

    df = df.copy()
    df[gene_col] = df[gene_col].astype(str).str.strip().str.upper()

    rows_html = []
    for sno, (_, row) in enumerate(df.iterrows(), start=1):
        gene = str(row.get(gene_col, '') or '')
        pct_raw = row.get(coverage_col, '')
        if pct_raw == 0 or pct_raw == '0' or pct_raw == 0.0:
            pct_str = '0'
        elif pct_raw != pct_raw:   # NaN
            pct_str = ''
        else:
            pct_str = str(pct_raw)
        rows_html.append(
            f'<tr><td>{sno}</td>'
            f'<td>{html_lib.escape(gene)}</td>'
            f'<td>{pct_str}</td></tr>'
        )

    table_rows      = '\n'.join(rows_html)
    safe_patient_js = patient_name.replace('\\', '\\\\').replace("'", "\\'")

    LOGO_B64 = "iVBORw0KGgoAAAANSUhEUgAAAT0AAABtCAYAAADTYrbeAAAABHNCSVQICAgIfAhkiAAAABl0RVh0U29mdHdhcmUAZ25vbWUtc2NyZWVuc2hvdO8Dvz4AACAASURBVHic7Z13eFRV+sc/507KpPeQRhJIQggJJRA6EnoXsQEWWNsKYvspytpB14ZtRd1FVFyWsrgqIIJKkRqkiEBooQUJJCG992Rm7u+PIYHJTJKZSSHK/TwPT5hz7z3nzJ2Z733Pe97zHiHLsoyCgoLCDYJ0vTugoKCg0JYooqegoHBDoYiegoLCDYUiegoKCjcUiugpKCjcUCiip6CgcEOhiJ6CgsINhc317oDCjYOYvcKoTP50xnXoicKNjGLpKSgo3FAolp6CgkK7o7CwkISEBIqKinBycmLIkCH4+Pi0SN2K6CkoKLQbysvLeffdd9m0aRMhISH4+fmRm5vL22+/zbBhw3j++edxd3dvVhuK6CkoKLQLSktLefLJJ8nNzWXx4sX07NkTIQQAp0+fZv78+cyePZvFixfj4eFhdTvXTfQ0Wh3pRRVczC/jUmE5VTVaVJJAJQnUNipCPZ3o4uuKm4Pt9eqigoJCG/LZZ59x8eJFli9fzpEjR1i7di1arRYXFxcmT57MP//5T2bOnMmiRYtYsGCB1e20meil5JWy/Vw2O5Kz+ep8NprMYjAnwYuzPT0C3Lk1yp9J0QH06ehZp/4KCgp/DsrLy1m/fj1Tp04lICCARYsWkZWVxYMPPsjKlSuZM2cO27dvZ+rUqSxZsoQnn3zSamuvVUUvr6yK5Qcv8I/9F0hNybWuktIqjp3N4tjZLF5dnwgu9tzbM5jHbwqnX4h3y3ZYQUHhunDy5EnKysqIj4+vK/P09KR3797s2rWLnJwcAOLj4/nkk084evQow4YNs6qtFhc9jVbH+uPpLN5/nm3H00Gra9kGSqpYueccK/ecIyTUm7dGd2NqbEdUkhJ9o6DwR6W6uhoAtVpdV7Zv3z4mTpxIeXk5H330EZIk4eDgYHC+NbSYUsiyzLIDv2M7fz13fLqTbYmpLS949biYksvdn+/G9fWNbD+b2aptKSgotB6BgYEAJCcn15WNHDmSbdu2MWjQID766COD4wEBAVa31SKit+FEOh5vbOT+f/8CuaUtUaVFlF8uYuQHW5n0RQJpBeVt3r6CgkLzCA0NpVevXqxZs4Zrk7mrVCqGDx/O2bNnKSoqYt26dYSHhxMdHW11W6I56eIv5Jbwl/8dJOF4utUdaHHsbVh2zwD+0q/T9e6JQj2UZWgKjXHgwAGeeOIJpk+fztSpUxFCEBAQQFVVFRcuXGD79u18+eWXLFy4kJEjR1rdjtWit/FkOjd/ngCVNVY33po8PqYbH07pjSQpM73tBUX0FJpi586dPPfcc4SHh3PbbbcREBBAXl4e3333HUeOHOHll1/mlltuaVYbVonem1tP8uLaI+aFnADYqvTnalrXx1efgdEB/PTQ0DaP9cstreLbxFR+Sclly8VcsosqQCeDnYoBfu4MCfFkdKQfoyP9LAq/SUzL54sDvxuUuatteX1iT0A/iZTwew4Jv+dQVKF/GIV6OtE/xJPYIE9sVS0/2VNRreG31HwS0wspqarBRhK4O9gxuJM30f6GkfMtKXpns4vZn5LHudwSKmq02Kskwn1ciA10J8bfHRsr3usz649wKrvEoOyh/p24tUdHAMqqNPx06jI/nsqgtFqDQBAb6MaQzj4MCPE2u82c0ko2nLjM/ou5pBRWYCMJXOxscLa3oYe/K+O6+tPF17XFQ7PKqzX8dimfjOIKtDoZJ3sbega4E+rlbHFda4+msj05y6AsNsCDBweGAfp7telUBkcvF1BarcFWJdHV15UBIV5E+bk1WndaWhpr165l69atdcvQhg8fzu23305YWJjFfa2PxaI3c9V+ViScM31QJRHXpQM3d/Uj0seFzl7OhHo54eOsn5HR6WTKqzWUVWs5k13MvpRcElLy+CElt9V8gVHhvhx+chRqW1Wr1H8tSZmFvLo5ia8PXjBP4H1deHdYJE8M7YKdTdP9++9vKdzzRYJhoauaY/83ipd+OsH3J9KhogHL28WeuUMieHZ4FB1c1abPMROtTse3iam8ufMMx5Kz9YJuCg9HpkX5M3tgGMMiOjRb9FILyvlw1xk++CUZSiobPtHZnof7deLxIRHEBJi/ZCn07Z+4WC+0av4tvfB1VrMy8RL7zmQ2/Ln6OPP5hB7c169Tg+JXWFHNnG8PsXrv+aYNBg9HHo4N5tVxMfi5Opj9HuqTU1rJPxPO8e7BC5RnNBAb62LPpK7+zB3ahWERHcyqd9bXB/ls+2mDspG9OvLkkAhe336aX882cq/8XPlgWCSPDI5ok99lfSwSvZd/PMbr3x81LLSRuLNvKHf26Mj4KH+c1dZZVUmZhXyccI5P9/8OZdZPR5tiZK+ObHk4vtWGulqdjje2JjF/w1GrrFnHADd23zeYPsFejZ5nUvRqrQFzP0ZHO/4zvR8zrfR5bj2dwZiV+y1+SN3UPdCk79cc0auq0fLshkQ+/vlUwwJrCiG4b2gE703uhZeTfZOnmxI97FRQrTW7SRt/NxIfHmpk5e5PyWXgpzuhsMLsumrb//ukHrw4Otoiy6+iWsPT64/w6a6zFn0nfTt6sOHeAU3GwJoSPWwky77/fq7su28wA0LbNt7WbNEz+sHZ2/B4fBdeGNWtWU+i+lTVaPnqyCUe33yCkvTCFqt35k0R/OeeAS1WXy3FlTXEL95J4plmhsxIgi//Moj7+3du8BSTomclj43uxse39zH7fK1OxyPf/MbnO860SPu1NCV6SZmF9FyyG01GkfWNeDiS+NgIegY2HsFvUvSsYOGdfZg3slvd68tFFQS+sRGKG7FOm+DxMd1YdGtvs4Tv14u59F+6B+oN1c1GEjw7Loa3J/Vs0FAwKXpWtrXqgSHcHRfa/LrMxKzg5COp+dyz7Bf9C0nw1JhuvDQ6Bk8nO6NzZZ0GcfkY5J1DLsqAkgy9FSKpECo7cPUDz07gEYrsHY6QDLtgb6viL/06cW9cCP/cc44nv0uE8uZbfssTzjGqSwdm9G25Wd2yKg19PtpG8u85TZ8sROPWmE7mgSv3uDHhayk+2ZqEh4Mtr03o0eS5Wp2OKf/+hY0HU1q9X9dyMqOQmH9sbZZYAFBQTq93NrH9sREMN3P4Zi2zRnTl2RFRBmW3LPul2e/h4y1JhHo48vTwqEbP252cTfzH26BKY31jOpl3fzzOxYJyVs8Y2LqTgTqZe77cg7uDLROiA1uvnWswS/SmffWr3mz1ceaXB4YwqJNhXitZ1iHO/YyctBH54hHkatMmvNFP3jMQMeIZ6HST0bkqSeKJoZFM7RXM1BX7WiQsZuZXBxnb1R9fl+b5tGq5Z9X+hgXPVc3fbopg1sAwgjycsFVJlFbWsDcll4U7zrD9WJqxCMrwwLK9dPFxYXBn63OHBQR7EdfBhWqtzKaLuZBXZvK8v284xvgofwZ2aryte1bub1jwOrjw3rBIRnXxw8vRHhlIyixi5/ls3t6bDPnWxU2mFZQ3LngCekR0YHSYD852NpRVa/jlUj77TmWYHgJXaRixeCen/zaOyA6NO9Ibw8bfjelhvlRrdeRVVLMtJbduyDqhTwj/uiPOwBo7mVHIb6cyTFQkMa5HR0ZF+ODpaI+rvS355dUkZRXx9alMLl/KM7pk7qaTPDqkC/YN+MF+vZjbuOC5OTCnbyiBbg6obSQKKmrYfC6bg6b6B3y97zz2NhLLrR0hCUHvSD+CXdWUVGvYdjbLtAGjk5n4n73kzp9slhuiuTQpev8+8Dvnzudw1+BwPp/aFyf7q5fIOi0c+xr54ArkwmyzGhTOHhA9HiJGgV/3qz6pBvBzdWDXnOHMXX+Ef2w6aVYbDVJWxf1f/coPfx3avHqAb45cYv2vF4wP2NvwxRWfWf3ZUme1LWO6+jOmqz8peaU8/O0hth65ZHi9LDNk2S+UvzQJBzvLVgmO6NmRhRO7E3eNb1Cnk/n+RDoPfn+E/LR67gJZZvyq/RS+dHODdX595BL/23veqNwzyJ1Vt/VhbJS/0ZCro4cjY6P8+fv47ny69zyPf30Qasz3iwHcuWKfacGzkZg/sQezBoXj72bsVskrq2LZr7/zzHeJxj/+8mp6L9lF0YuTLJ/ddXdg+4M3GVmKWp2OH5My+OlUBv+YEmtkFX2+z3C2HaBvlD+b/zoUD0fjkRLAB1Nk/vPrBe5fud/wvhVV8O8DF5g9JNzomopqjX5Ia0Lwwjv78OqIrkyNDTZ6338HTmUW8f7OMyz9Jdnoc1qRcI6JUf5M6x1isq8N8ejoKOYNjyLY06murLxaw7IDF3j0+yNQUmV4QXElc9cnsuzu/ha1Yw2NfvLFlTU8sO4wi+7ux39nDDQQPHLOwKp7kLe+D2YIngjphXTLmzDrJ8TQpxH+Pcx2zAoh+GBKb768f3CTItkUPx66yJbTpp9s5lKt0TJ19QHjA052HHpmLA8ODGsyPCTUy5mf/noTk/qGGh/MLuGNrZYJ/NSBYWydFW8geACSJJjSI4hLz44nKtzX6LqitEK2NeCPLCivZtqq/UblfsGeJD89lnHdAhr9DG1UEo/dFMGx58eDu/l+36X7zrM/6bLxAW9nDs4bx4Lx3U0KHoCXkz1zh0eR9PwEHAOMLbryy0X8Y6eFvihXNZkvTDI5NFZJEjfHBPKvO+NMWmBH6vsi3RzY8nDDggf67/t9/TszJ76L0bHvT5m4L8Dc9YkmfXhTB4Zx5pmx3B0X2qDQR/m58cX0fux6chTYGz9op//3ALmlVSauNM2iu/vxye1xBoIH4Ghnw5ybIkieNx5MfH7/2Xee/BaexDRFo7/MxXvO8eWtvXliaKThgUPL0S2fiZyZbPrCaxCd+iLdvxox9QvoMsbIh2cJ9/fvzKf3Nn8y4vlmWoz/PXTR2Aqxkdg9Zzi9O3qaXY9Kkvju/sFM6GP8FH1j11lqzFy7HBTi1aTvxcnehl2z4/Uxk/X4dJ+xJQewZO85qPdl9wv2JOn/Rjf6o61P9wAPXh0W2fSJV3ho8wnjQmd7MuaNMxL1hojycyNj3njwcjI6Nu+H41RZYHmuntbP6jCf1FLD70m4lzPuDubdu3kjovi/sdEG/8Z2MX5w5ZZWsXiX8QTT1IFhFvnkhob7suvxkcbCV1rFot3mTWA9OCzSWC/qEebjwqb7BhsfqNHy38MXzWqnOTQqegNCvQyc6rKsg21voNv+Eeia+NLY2CHd8ibijsXgHdEinQWYNTicV2+NbVYdh09n8JsJn4m5vLHrrFHZ9P6duCnM+AvZFCpJYtld/fTT/ddSXMm3ialm1fH3EV3N+mL7OKu5y0SoyrcXjGcsdTqZ53cav8+dDw6xSPBqyTbTUthyOgMyi43K18wYaHGUgKvall33DzEeHZRX881R8+5tdEQHpvUOtqjda3Gr56JIziqmotq8SYYQTyf+cWtvg39Pxnc1Ou9fv5wzChUJ6+Rj1STE0HBfVs0YaFT++u6zaMx4CP9tpHH/TDE2yh8bf2NLPMGcScFm0qjoxYdfNedlnRa+fxrd4XVN1+rgjDRtMXQZ0+wOmuKVsTEMiWneTM/8zdZZe2VVGpIvGH8wzw4z78M2hY+zmjtMTNn/2ICD2QBJcLcJS7EhHupvYvY6u5jKepZPYnqBUSxev24BzZoEMIfV9X2cwLCeQdzWs6NV9Q0N9+XewcY+sKUHTfhjTfBo/07NWhlxU0g9y7+sisfXHrbI0myKD0y8lwUjzXsQmmJa72DwcDQsLKxgx7ks0xdcoVMnbyJ8XM1u50UTD+CdaflmX28t5o81N89HPrunydOEmy/izn+BR2gzutU0X88YSMBrG6DMfF/Dtfx45BIF5dUWWy0HL+UZTUNHhvlaNKw1xbzhXfl2v6HTe8tFM2LGnOzMWs1RSwdTM9cylFTVGETH708xtoSft2CIai0/m2jXQ23LR2YOr0xhozL+8e88b55FEWPCGrGECV39+XhLkkHZ0p1nWJp4iacHhBHi6Yi72g43B1viOnoS6O7YQE2mqajWUJRu7DecFmu9daqSJF4ZGsFr6w0XIuy7mMforv4NXhdioQvAz8XYcs9ug7X85oner18in9jU9Hn2jjB9Kbg2fGNaCn83B/51ayxzVho72s1CJ7PhRLrFKxMS0wuMyiZHNj/2q2+IF6htoPLq0CfbnOBsC60QVQNPf229MI8D9Yf/NhKTm2ldN4VWpyPNxP1dd+AC6w6YZ5mZTWkVOaWVdUskW4uxUf7ERvpxpP5kUWEFH2wy4bv0dua2cF9GhfsSH+5DVAe3Ri3NQ2n5RqFPs+NCmr3O+qEB4SZEr3ErTNUS30VLVtxYSdN35vfd6HYtNq+ykXMRbSB4tTw0MMzYDLeA/5np17mWEhMhAd4mgrStor7VqZOp1rTcMMgSsuvPojnatXrGmpIqTZsmpbBkRtJahBAsmxoH5ia9yC1l7f7fmbNyP9ELNjDo4+2cyzH2cdaSUWQc1hPkZv1vopYAN+OHwcViC5fQtVMaFT1ZU4m86TVMhBUbISIGQXTzUr5Yiq1K4v0x3Zo+sQF+PJlu9gxpLZKJp1lVS/1QTQxTNW3w5DOFrv5n3gYpuqzP7GgdGl3bCGyPQA+OPTvWorCdWvYnXabLqxtY2sAMe30LHUBt2/xsOipJMvrMq7XX57vY0jR6d8SR1chlZgyx7B0RYxa0UJcsY86QLmBt1pBKjcWzuH4mfGIZzV0m1Q7xrG+ZtEH8lLO9DbRh+sO2FNnuAR5kvDCRh0d0NRkL1ygaHQ8t32cyvtRZbVxXS8S6FVfWGA01PUy09Uek4XdRXY7uwHKzKpG6jQbH5jnyrUVtq2JOv0786+dTVl1/NL2oyWVY19I32Ph9Lj6Uwke39bYqh1t7pW+QB19dW1CjZee5LLNTD1mDrUrCxs/NKLnAoJhAuvu6tHh7re3Pq4+fqwNLpvblo1t7s+1sJgcu5XEhv5zk/DKOFpRRXqXRL9MylR5MlrlnzSFyXpxkUNzD3zh11qbkHN5oZl93JRsvOIgzEez9R6Rh0Tu8HCrMzNLQ/bYW6o51jIzoYLXoHU63bIo82t/NON1QYQXfJqYy3YLQkfZOfxOphRbuONOqogcwJcSLb+uJXmd3Rz6d2rdV221L7G1VTIgONLnAXpZljqQV8PxPx9ly2DB8Jze1gCOp+cReEykQ7OkELmqD/IKHT2eQlFlINz/zcwnW560dxqtWBv5Jtlw1bZpUFaP79b/m1eAbCh2s36SjJRgWbv0P8YCF6atUksQtvYzDAV4z8SX5I9MvxMvIbbAp8RLpha278dLkaONdrlbuO8/h1NaP32oPCCHo3dGTTX8dyi0mIgs2mlied0uU8eThe81IAXY2u5h9J+u1IwlGdmndB15bYVr0zm6BKvO+3FLUhJbsj1V4OtnhEmjdU+1YpuV52uYNN45XO5WczeyvD1rVh/aIrUpi3k31VtLoZKau3G/VjLKnmfGQ03uHGPtotTr6f7nHKIC6KbQ63XWb/Qb9xlkrDl4w+LfGzFU2QghmDzBOMZZq4qHz9FDjNbr/3nue3SaGqE1RVaNlxlfG3+PRPTtaHEPYXjEpenJaovk1+LTcErPmcKe1w67SKizdJmRQJx+6m3jqLdl++k8lfE8OjdQP5a9h74l0Rn66yyIxySiq4MND5q2ptFVJvGwiCFqTUcTQf+0gp9S8SaPaHIC9Pvy5TUJTTFFUWcPMpXsM/t3x6U52JTe+sqGWs7nG7qUUE5mXh4b7ElB/TXKNlviPt1kkfFU1WkYs2cWvJqzJF0ZYv+KovWFa9NKPmV+De/vwY422VvR0MvlWJCldM2OAkSCAXvhu+mQ7W5uZyaU9EODmwIe39zYq33Minfh/7eRMVtNW8uZTGQS8sdGiLNgvjYnGzoTT/OCpDHxf38impMuNPqi+P55G8Nub2HgwhVPJ2fi8/SMnM1ouC7e59AjwMJnwYNSqA03uz5xVXMmTPx43Kg9tIOxl/b39jcOKqjTEf7yNlb+loG0iPOdsdjFDF+9k7wnjvJV3DOjc6r7ctsR4IqOiEApMp68xQkjIbgFtGWXQIFF+5q/5q09RRbXFyQsjfFz519S+JleE7DmRzpgT6eDvymsDwgj2cMLNnL1DmpPttpV4YmgkXx1NN0r1tD/pMl3nf0+/bgE8PyyS0ZF+danHzuUUszs5h4V7znHOzOVe12Jno+KX+wbTd+EmqB9HWVjB+I+2gacjj/UJJT7MBxd7W0qqajhwMY//nswwTsCZW0rMwk38+PBQxncz9hm2FpIkWHZLLPd9abh8U5NRRMfXvue5kVHc1Mmbzt7OeDnZk1pQzu95pRxJK+DNXWeMc84B9/YJNdlWXLAXfxsfw8If6glllYYZXyQww8ORV4ZG8NCAcALc1KgkfVLbHcnZvLXjtLEPrxYPRz6f9ueZRAJTonf5iNkXCxevZqWKakl8mxF+UGzler9HhkRQUFHNi2sOmz4ho5hX1pl/P9sjQgh+nhVP/0+2c9LEgvNfky5za60g1loaLRBQHRfsxXez45ny6S5j4QPIL+eTrUl8stXMCitrON9ABunWZEZcKC9vP01q/b03Kmp4e+Mx3ragrsgwX4Y0klH7zYk9uZBfztemApkLynlt/dGrS8sk0fTn5KLmxBMjzU6F9UfBaHgrZ5pYD9gAsrb1A1bNxcfZ+jTTxZXWW1gvjI7mw7v6GqeG+hPhZG/DgcdGMLR7UOMn6uQWXTt5S/cg1s+O169Jbg72Nix7YAiP1Z+YaQMkSbBhxgBwaV4adN+OHux/bHijSwElSbB6xkBmmPM+m/qc3B048fRoo13d/gwY/1JrLFhdUF6iz7HXDrBRSSYTZJpDRTPT/DwZ35UTL04kKMS8BJd/RJzsbdj16HA+nTHQKhEylcXYHCZ3DyLllcnERvpZdb1nkDtJz0/gL1ZuedkS9Az00GcLtjLAenRsMCeeHGWWxSVJguX3DGD1X4eClftN3DGgMwXzJ/8pBQ/MSTjQGLIOUdH2DuIGsdK52JST1xyi/d259Nx4vn9sBINaORvJ9WTW4HDy3rid12/rbdJJb4CNRK9IP5beN5iSRvbhaIoQTycO/d8o1s0ZTj8zfXJhnXxY88gwcl+cRJTf9V9JEObjQt7fJvDCpB5mW33du3Rg8/+NYsuseItXj0zvE0L2q5N58eYeJlOzGyEEo3p1ZPvcMXxz3+A/3ZD2WprvkCvLuW5L0IywcmjVUuuohRDcHBPIzTGB5JdVc/BSHvtSckkrqqBaq7NqrWdzUwS1Bp5Odrw4JpoXRnfjTHYx+1PySEwvpKiqBltJ4Ka25abOPoyK9MPRws2NGkII/V4fU3oEcTG/jN3ns9l/MY+T2SVUaHQ42Eh09nAkPsyXmzp709nbcquqv78b2nofkoul62QbwdPJjjcm9eTlsTFsPJnOkbQCfrtcSH5FDWqVhLOdDXFB7gwK9aZfiFezdwbzcVbz+sSeLBjXne3nstiXksveS/lcLKqkWqfDw96WvgFuDAjxYlQXP4KakbHoj4TRZt/yzneRD/7P/ApGPY2IvbvFO2YNYs5Kq4Tvu0eHc0tT/qp2QGllDZfrpfexlQSdLPiB12h1XMgrNSoP83bWZ9ZoRcTsFUZlTW32rdA+ySmtpKBeqJeLvW2DGzaZoriyhsx632c7lUSol3OL9LEhmv8YO7MV2onoWUtbpzSyFme1LV3MCX1pBFuVRBdf68N7moMicH8efJzVzU7Y4Kq2xbWZ32drMH6021hmUsupJ6DMjLTmbUEzt4dUUFD482MkesLivS1k5DObW6Y3zUXRPAUFhSYwjtPzDLW8ll+XQ3XbB34aoVh6CgoKTWDk0xNeYcgIzEkRX4tckodI+AeMfKkl+9Y+Kc1B983D17sXCgp/WqSYm6HvA61Wv/FEhp0TeAdBrmWb5ugOr0eKmgQBvVqqb5bTBvs4oK2y+N4oKChYQIl5WWisxWSMghRojXDJ6DYtaB/DXAUFBYUGMB2y0m0SHN1geW15acjfzkbc8aneYmxrrrdLz7cTwuaaSHZZvz5ZlBchlzae+Vd69re6/+vejWutHv6p+cPcQ1s1IrAb2KqRizIgu3l7+gq/cBAq5OJMKLM8Ka7Z7bj5gKMnck0F5F5q+oJ2imnRC+oDHgHmp5i6Bjn9FHzzMOLWT8DRo7n9s4zrPJEh3fIPcDcMcq7tkcg5i/zbf5BPtJOZ7j8aDi4IN/36Wznz3HXujPVI8Y8g95mBUOkfjgKgMA154zzkjLNW1SmmLdUbGRYuLLC4nf73Qc9piKwkdMtntlo7rU2DIfhS98lWVypfPoO88m64uM/qOqyivczeluVC1kn9v7zz+qQMPl0Q499AGvbo9e7dHxKpyzDEjNWIGauvd1esRnQdDv0eBMkWOfEr5F3vQUGK/kF5ywcg2t+Swz8jDa/IiJkCez4HnXUZSOSiHOSvH0dEjUAMexaczd9m8Q/PyfXodi2++trRFTHuZQgbDn3vhxMbjIYH7Xo49geh3d/DqHEAiLNb0G19T1+WdgRxzyqEix/4RyBftn5DHwXzaFj0nLyRYm9Fd+jbZjUgn9qOfHY3InwgImYKcugQhGRdCqgmaS+WXn3Ki9Ftfh3xSDxCSIhOA5DriV5T/ijh4gVdRyF8IkFSQXk+8tltyGmN5D/0CUX0uh3h4An5v6M7uBqp712gskNO3ql3RVzbh543g3sQctZp5KwziN5TEY6eyLnJcHwDcmlBg02JoGhEl1H65BPaGshKQndyU8MbTNnYIiKHIYL66Idmmkrk/AvIF/ZCdoph3cE9IGLk1X7GP4KcvAM53XAHuiZ9eg7OSNHjwSdSv/KoLBcuHUCXfABTIVoisCsifAS4+AICSjIhPRHdhYOgtTwHo7DV+7nl8qsrmOScC7DhGf3/izItrrNZqJ2Quo4Evxi9f7GyGHFhD7pk42zg9RHOHog+d4FLByjOQD7yNXJJUaIqDAAAGHFJREFUA35rJzekqLHge+W+VxZB9hnk5IRGv1OtReNrb4c8hjj9M3JZM9NHaTXIZxKQzySAvSP4hiF8u4B3BKhdwUat/yELleFfSaU3+SXVFYtT1g8VZRkh60CnA3SAQFbZsvru/lRasftVXMc2yBJTVoSoKgG1G8LRy4IoSBChveH2xUYPC9FnBuLYN+g2LzR9zW3/RKiurm2Uuo5Hdg1AqOwQpTlGoke3iRAUh8g5C+7BCFv92koB0HsGrLwHuch4oxlp3AvGex/HTEEa+Ai6b/5qJGLYqZHu/Q94hRn2GRBDn4akDeh+ePXqgYCe0GnI1df9HkSUZBqJXmOI4B6IKR+Cfb11x31mIDKOIa+epRfr2vd000MwYLbJuqSyXOTld1n+g809CyEDEJ3ikcUikHVQU418eqdl9bQEHgFIM1cbTDgKgF7TkS7uR/ftkw2O8mQnb8Rfvjbw2YteU5FX3Gs0DyA6dEZMXwZ2JjK4jNQitr6KfOzHFnhD5tO46Nk5I4Y/hbxxfsu1WFWOnHocOdV40xNLMCUa057eZ/Ajb79YZpHKKYcRCR8iu/pB8k6oLkeExcOAh6HHnYgTG5DTkwwvuulx/b2oLELe8S6U5SC632beMkOfLohDK5CzkhC+XZH7zEQ4eiD6TEfe/pHBqVJY/6uCd34H8tmfwdETMeCv4OiBGDsfecX9htfETNALXnkBcsI/oChd/zAI6Q8hg9Dt/czw/Z/5GZx9rmbzWfcYco6JlOgNIalgwpt6wassQv71SyjNQnhHQPQUOLzSQPBQOyH3+6s+RP/wSvg9ARDQIQrRbRIc/doqC0V3YgNSn5ngHoQ0cCa6vcssrqPFKLiMvHkBImwonN6EXF6I8I9GHvYsImQAInoM8vGfTF4qnH2v3hfvcBj6FMLeFWngg+h+/LvhyYNm6QUv8wTygS+gqgRc/SF8OMLeFV3Sz23wZg1pOstK1ETEsTXIlyzYIU2hxZF/NXTgy5fPIIXFg08kokMXQ9ETEsKvu/68/UuQT2zS/z/1GCJiVNPuhQsJ6LYv0l9zciuSix9EjtXP6tcnfJj+b2EqunV/01svABWFiHF/R/h1R3ZyMwylUF+xtsqykdOOQlHWldHAbtP9KchA5FydsTVn+HUtIrCb3mcGyBueQU7R71si8zMkfGZk0Qg7x6v36NKvyJdP6YfpF35D3m+cHsssvDsibl1U91Ie9CiiJKtOWKQRTyD3nIq4uB/d2mesa8NC5NM7kE/vuPo64wxSUG+IHAe+XQHTokdhGvK2D/X/v3AIyasTdL/D5MIEob6SwDU3GS4n6UeN8lE4vsmi0U5LYlZqKTFxIay8F7kkr+mTFVoeGztE7K2I4H6gdkeu9V26XQmPkeyMzq+j9JoMOJoavUXTlOhV1ov1Kr0ypDWVgcfuSu6z4vSrggdQeM1WgnZOBqInX9wPAx9B+EQiPbheX1hVrLf4Lv2Kbu9SszebNwv11XyDcnY9C9HEEE4uydMP8X26IKZ8pLf4tNWI0mzIPIl8aIVFQ2sAacIb4BoAqQeRc88hYu9GHvsakq0a3eF14NwBYaOG8sbjOVsMISG6j0N0GgLOHa75TnXUH1bZNixKVcWGr0v03w/Z1jiXnvx7AiIoDmKmIGKm6F1TFQWQnwJntzR7zsAazMun5+yDuP0j5P8+CNUW7KGhcBUhIavs9QNbjfGGzY0hTX4DwoYja2sQJZl6fyYgSyrTA+VWm8+xpOKGn+Ny+mnEV/chd5uI8A4HJx9w8gbfKPCNQvKNRPe/x5rfXWuRdei+no3oOQUR0BNc/BGOXvqHjFsQhA2DZXdAgZl7G7v7Qodu+qp3vIuc9TuU5SKGPAEjX0RSu0HIAP25FmzM1Ryk+NnQ9wFkWUYUpyNqxV/VsrsbygdWQUkmovNNyO7B4OiF7OyDCOoNQb0Rtg7WW89WYv479IlEmvwWujVzDZ/oCmYhokfVTQzIeSnmX6h20oe6APzvfnTXWBjS3Z9DYKzxNZqr/imhdrkqP5Kqxb/UdcsOnX0Ny68JURI1FUYSKGcmQ/p7BuUiagRi0jvIHfuBhUkvGqXqaqZo4RGAXH6tpdJAO1UVyPuWGx5xdEXM/K9+qBzYHdlc0RNX73ntRlryvuWIyiLkkS8hBusFXq6pRD6707w6m0v0FP3fTS+ju+L+gAYmpepTz6KT1W76SShNlclPTD6zGzlp29UCIcGopxG9piNC+rdj0QPodBPSuBfQbXpTEb7G8I5AxFyJybK1h6A+yF3G6o/lnUc+s8v8umT9jLUQkn54dEX0hKs3sksH07aXTgu558A7Arn/w4iCVOSyPESvO1t8n2L54n5EjzvAszPS2OeQz2zRT2DEX/FL5SYbOf1FaCxi8vvISRsgeRdySRbC1lE/kQGI6jLkej8fWVdT916lXpPRpfwKheaFeMiXTyHK8/XhNDe/h0hYBKXZ+vCf2LuRf/478oVDVy+wc0C66zOoLEY++T1y1mn9sMy/G6J29rfKOOV+gxTnQFkeOHkhjV2AvOsDKMpAl34UKeOYfnYaEBmJ9QTZMoR/d4hpYBlaUbrh5GGtZefagTrhd3AGTzN2jfPshOh/r37o6tUZUSuSWUlGp0qDH0Dufjsc/R9cPIBcVYpw8kL4dgVAriox/w22EJb/AmKmIDn7ovv+uZb1u/yZ6DwU0XmoQZEA5Ixj8N0zlj0wqsoRZ7dC5FjEpHcQY8oBGeyckBsJHJf3LkZM/gDh0gHu/EwvGOUF+iFyC85wy6d3IaJ3Qudh0OMOvQDWHquphK31ZvMQiIGzwN4VEXsPxN5jLNwJHxq3k3YUdFr9BMPoV5C2vaH3hZmDpgZ5ywK94Ln4wYS3DHs08nnkf99VN4MrRY1E9u6CkFSI4P7G/buciGzJZIq2Bnnb6zDpPYRfDGLal/p2a9+btgZkLSJ4ANL4F9Ftess6oyJyHCJynOljp380FL0Ta2DAbMTgx6H/XxE6DbKtk/4B21Q7FYUw+FGkof9XVyRrqpD3f2FwmnDxhJ5TEU7eMOQJGFLPQVJTAfsMr2kLrHvshw5CuuffyGseNxm3dcOSvM14ZzhNFVQUIJ9PMA4rMRPdDwuQ0g8jB/dHqN2hPBf53DaEe0fwCEHOM16wLp/ZDf97AHrfhXD0huJ0dL98hrjvmyuV1hhfk7IXUZwB6UcMyzNPIJI2XJ3QMDyKbs0ziMh4ROQYcPICTbX+qZ/4LXJx/a0EZHSrZ+vj5joPAa9w/WSItgo577w+fMLUJEFBBnz7MHSbhGzvDIVpZt69K62e2wtfTkH0vONKkKz6anDyyU0GISu6oxsRF/YjR47UW09OPiBroSwP+cIe5FPbLRYl+UwC5NyG6Hkb+HQBlR2U5yFfPgpntoGLN/SaBpINImIg8tlfzK/89E9Nb/NwOdHgpS5hKaLgEqJzPMLJB6qLERf26t9nQC/9A7r+e7h8HGHrqJ+0Op8AcTP13/fCVOTfVhilXJNL8pGX3IyIGIIIjgP3EFDZ6l0imceRj29ELs4x/322EEa7oVlEeQFsftniEILWQrRFnF5RGrrPprRuGy2Fhz8UZdcNZUSHzoiZXwMgf/ck8jkLflgKCm2E1Od2GPF8q9XfPAePowfc+gnS+Z3I2xYiF7W9aiuYRnQZgpjwlt46+3032Kqh6wT9wcJU5PPt40GloNDWtIxXO2yYfnnN/s+QD642mD1UuE5UFiOX5yM8QqDPNVsvZp1E9+PLVieSUFD4o9NiU3nCRg1DnkDE/QU5aQPy0bV/6ESDf3TkS8fg81shoCu4+oFOi1xwqdkJKxUU/ui0cNAW+jWUve9F9L5X7zxN2oB8+QRyTopiXbQ1sk4/eWLlBIqCwp+R5k1kWICsrYacM/qI89IcRHU5aKtBltHnVZf1sVm1r/Xdu/Lnyl9Z1s8uNcTYv7de2ipqu6CDUsV3qaDQWgh7p6vLG1uj/rYSPQUFBYX2gJKfWkFB4YZCET0FBYUbCkX0FBQUbigU0VNo9zTmdm5tl3RNTQ1lZcYb2BcXF7dI25WVlZw6dcqoPCcnh8LCq9s0aDQaqqurLapbo9Fw4ULzQpSSk5MN+tEUFRUVZGe376WpiugpXBfKy8uZN28ezz33HBpNw5vs7Nmzh2effdbksYSEBObNm9daXQRg69at3H///Ubl8+bN4+LFi82uf8mSJWzcuNGgrKCggOnTp1NQcDU7zfTp00lKsiz06O2332bt2rVW9626upr77rvPoB+NUVRUxKOPPmryIdGeUERPoc3RarXMnTuXsLAwvL29Wbp0aYPnnj59Gi8vL4uPtRRqtbrBH7EpC81Sjh8/TkREhEHZyZMnqaiooGNHfRbj8vJyUlJS8PGxbBvVPXv20KVLF6v7dvjwYdzd3QkNDTXr/DfeeIMxY8bQqZMZ6amuI4roKbQ53377LaWlpTz00ENMnTqVdevWodOZzlqSmpqKv7+/yWOXLl1q8FhLceHCBby9vU0eKy21IKeeCbRaLadPnyY8PNygPCUlhbCwMGxs9GsHMjP1eQMb6ocpysrKyM3NNRJUS9izZw8DBgxAmLG1alpaGsePH+eOO+5o8tzrjSJ6Cm1KVVUVn3/+OXPnzkWlUtGxY8cG/VrQuLA1Jogtxfnz5wkKCjIq12g0zR7GZWVlUVlZaVR/RkaGQVlmZiaurq7Y2zeRPqpeHYDJvpuDLMvs2bOHQYMGmXX+jh076N+/f51Qt2cU0VNoU7Zu3Yq7uzu9eul3zhJCEBIS0qB/LDMzs0Fhu/bY8ePH2bJlS4v399KlSwQHBxuVl5aWNnsiIz09HbVajZubm0F5RkYGfn5+da8zMzMtHtpmZmbi7OyMs7N1KxtSU1NJT0+nb9++Jo9nZGSwevXquntw+PBhevToYVVbbU37l2WFPxVbt25lyJAhBmXOzs4NzvgVFxcbiYKpYxkZGc2eqayPLMukpqYihGDx4sWoVCrc3d2Jjo6msLAQlap5Sx7T09Px8/OrGz5qNBry8vJIS0sjNDSUS5cu4eDgQEZGhkWiV1hYyJkzZ3Bzc+PChQuo1Wo8PDxQq9Vm13Ho0CGioqJwcXExeTw/P5/Tp68me629TytXriQoKIghQ4a0W6uvffZK4U/LyZMnueuuuwzKnJ2d6/xjhYWFvPPOO2RlZREREUFVVRWOjo6AflbznXfeITs7m4iICKqrq+uOjRkzpq6+9evXk5mZyaxZs5rV15KSEsrKyoiOjiYwMJCSkhKSk5P55ptvyM3NxcnJqe7czMxMvv/+e7Kzs/H29mbWrFkIIZBlmZ07d7Jq1SoqKyuZP39+nZ8tPT29zlLdtWsXL774IhqNBo1Gg6OjIzU1NURFRZGWlmaW6FVXV3P33XeTkpKCjY0Nrq6uLFu2jMjISPr160dmZiYLFy7E2dmZ+++/3+Ce1efw4cPExcU1eDw6OppXX30VgMTERNLS0ti6dSuPP/44UVFRVFZW8sMPP3D8+HFkWWbOnDmtPulkLsrwVqHNyM/Pp7CwEG9vb44cOcL333/P8uXLuXjxIhUV+m0x33//fWxsbFiyZAnjxun3e6i1GN5//33s7e357LPP6n6wtccWL17MV199BcDRo0cpKtJvkHP8+HFSUlLq+iDLMsuXL+fpp5/m7bffZt26hvfZyMnRJ5aIjIwkNDSU7t27c+utt/Lcc88B1P2Ijx49yuzZs+nXrx+PP/44K1asQKvVJ8ZYu3YtCxcu5JFHHiEyMpJVq1bV1Z+ZmVk3jA0KCmLp0qXs3q3f8Pxvf/sbc+fOZcKECeTm5polelqtlnnz5vHzzz8zadIk4uPjefXVV7n77rsJCgri+eefZ9asWdx7773Mnz+f3377rcG6jh8/TmysiZ32rrBx40befffduveh0Wh44YUXiIqKoqSkhIceeojS0lKee+45Ll68aGAVXm8US0+hzbh8+TKgjzkLDg4mODgYb29vZFlGlmUKCwvZsmULX331FTY2NnUTBRqNhvz8fLZs2cLXX3+NSqWivLy87hjoh7e1Zfn5+cTExPDpp5/y3XffUVZWxooVKwgNDeWHH35g06ZNrFixgpqaGuLj4xk5ciSurq5G/a0NBvbw8DAoz8/Xb8hdG8rx4Ycf8tBDD+Ht7c3cuXN54okn6vq/aNEi3nrrLfr06UNWVhYJCQl19VRUVNQJZ1hYmMH7ubbNoqIifH19KSgo4OWXXyY0NJRnnnnGqL8ODg7069evrh5f36vbcv722294eXkxceJEhBAUFBSwdOlSk9ZcWVkZaWlpREVFGR2rpbCwkOTkZIC6Nmvv4Zo1awgICOD2229n0aJFuLq6NugbvB4ooqfQZtT+yNetW1c3LAV45ZVXcHR0JDU1FQcHhzoxqRWIkpISSkpKcHJyIiQkxOBYaWkp3t7eVFdX4+Cg34+1urqavLw8duzYwapVq1i5ciW7du0iNDSUNWvWcOedd6JSqZAkCVtbWyTJ9IDH1la/30p9311KSgqurq4EBwdTWFjI8ePHGTZsGF988QVPPPFEnUM/ISEBDw8PBg8eDICPjw+BgYFoNBpsbGzQarVGfq/atq69P2VlZTg7O/PUU08RHBzMvn37SExMrJsMaqjv1w6/T5w4QWxsbJ3/MCYmhmXLlpm8NicnB7Va3WiITE1NjcH9BrCzs6t732FhYbz55pvExcXx1FNPNXiPrwftpycKNwSOjo4GP2jQWzLe3t51IlNr9W3evJnhw4eTnZ1dF65xrUVYewxAp9PVTWrIssyGDRu455578PLyIiQkhKKiIrRaLSdOnKibjU1NTSU6OrrBGc7AwMC6Pl3LyZMn6dWrF5Ik1YnU2LFjWbBgQZ3gaTQaCgsLkSQJWZY5fPgwzz//PImJiXz00UeAoS+zFiEEarXaQAzLysrYsWNHnaV3xx13cPbs2Ubvc/06SktLDUJefv75Z7p27Wry2srKyjpBq0Wj0VBVVVX3WqvV1t3vWiGt/WtjY0NYWBivvfYakydPrrsHja28aUsUS0+hzXB1dTX44dRSWlqKp6cn4eHhBAQE8PHHH/P7779zyy230LdvXzIyMoiLi6NDhw588sknJCcnM2XKFPr06VM3ZLa1ta0bztnb2+Pn58dtt+k3oXZ2dsbR0RFZlrGzs2Pp0qUcPHgQtVrNO++802B/HRwc6N+/P6WlpQbCeOLECSZNmgSAi4sLw4cP56WXXmLu3LnY2tqSlJREVVUVsbGxvPfeezzyyCOcPHmSV155hdGjR7NkyRJkWSY2NpZly5aRnJzMsWPHGDFiBO7u7oSHhxsEa6tUKnbv3s0LL7yAra0tsbGx7Ny5s9F7HRERQXHx1Y3Do6Ki+PDDD+nVqxcHDx5k586dDa6EUalUlJaWsmvXLvbv38+ePXvIysrC39+fTp06MW3aNGxtbev8jLUCqdFosLe35+abb+bDDz/Ez8+P0NBQ0tPTOXToEA8++CA//PADarWasWPHkpaWxtq1a3nggQdwcnJi1apVdO7c2ezYQGtRLViwYEGrtqCgcAVPT0/i4uKMhk2DBg0iKioKW1tbxo4dS1FREd27d2fatGkEBwdTUVFBYGAg48aNo7CwkB49etT5BSsrK/H39yciIoLw8HDs7e2JiIjgzjvvrLNEqqur6dChAwEBAUycOJGhQ4fSu3dvVq5cSU1NDTExMQ32OS4urm42tZaysjImTZpUZ9kMHTqU7OxsNm/ezPHjx4mJiWH8+PF4e3vTsWNHtFotzz77bJ3vKy4uDiEEnTp1IiMjg19++YUhQ4bULd/Kzc2la9eudSEmbm5uxMTEMHXqVIQQeHh4kJOT0+gSMy8vL86ePUvPnj0BvQg6OTmRkJCAi4sL8+fPNxl/CHp/YkpKCtu3byciIoIHHniAZ599lnvvvZdx48bRsWNHgoODiYiIwMXFBXt7ezw8PAgKCsLR0ZEuXbrg6+vL5s2b2b9/P3Z2dsycORMHBwc+++wzHB0diY2N5fDhw/z000+MGTMGW1tbPv74Y4KCghq0QFsKJXOywg1LYmIiTz31FFu3bm1XMWWFhYWo1WqL4upMce3ssMJVFJ+ewg1JVVUVQUFBdbF37Ql3d/dmCx6gCF4DKKKncMNRXV3NK6+8UrdMKzc393p3SaENUURP4YZj3bp1BAYGolarsbe3b/ZyMoU/ForoKdxwHDlyhMjISGRZpqKiwupMJAp/TNqP91ZBoY3o2rUrBw8erJuBrE3WqXBjoMzeKtxwaLVavvrqK44dO8acOXPqVnko3BgooqegoHBDofj0FBQUbigU0VNQULihUERPQUHhhkIRPQUFhRsKRfQUFBRuKBTRU1BQuKH4f9V9+sywca1JAAAAAElFTkSuQmCC"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width,initial-scale=1.0">
  <title>Gene Coverage Metrics</title>
  <style>
    body {{ margin: 0; font-family: Arial, sans-serif; }}
    .header-wrapper {{
      background-color: white; position: relative; z-index: 2; height: 150px;
      box-sizing: border-box; display: flex; align-items: center; padding: 0 30px;
    }}
    .logo img {{ width: 200px; height: auto; margin-right: 10px; }}
    .logotext {{ color: black; font-weight: bold; font-size: 18px; }}
    .container {{ text-align: center; margin-top: 20px; }}
    .controls {{
      margin: 20px auto; max-width: 800px; display: flex; gap: 10px;
      justify-content: center; align-items: center; flex-wrap: wrap;
    }}
    input[type="text"] {{
      padding: 8px 12px; font-size: 14px; border: 1px solid #ccc;
      border-radius: 4px; min-width: 300px;
    }}
    button {{
      padding: 8px 16px; background-color: orangered; color: white;
      border: none; border-radius: 4px; cursor: pointer; font-size: 14px;
    }}
    button:hover {{ background-color: #cc5500; }}
    .table-wrapper {{
      max-height: 700px; overflow-y: auto; border: 1px solid #ccc;
      margin: 0 auto; width: 60%;
    }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #000000; padding: 8px; font-weight: 100; font-size: 14px; text-align: center; }}
    th {{ background-color: orangered; color: white; position: sticky; top: 0; cursor: pointer; }}
    @media only screen and (max-width: 600px) {{
      .table-wrapper {{ width: 95%; }}
      table {{ font-size: 12px; }}
    }}
  </style>
</head>
<body>
  <div class="header-wrapper">
    <div class="logo">
      <img src="data:image/png;base64,{LOGO_B64}" alt="Anderson Diagnostics Logo">
    </div>
    <div class="logotext">Gene Coverage Metrics</div>
  </div>
  <div class="container">
    <div class="controls">
      <input type="text" id="searchInput" placeholder="Search by Gene..." />
      <button onclick="sortTable(true)">Sort Ascending</button>
      <button onclick="sortTable(false)">Sort Descending</button>
      <button id="downloadBtn">Download CSV</button>
    </div>
    <div class="table-wrapper">
      <table id="dataTable">
        <thead>
          <tr><th>S.No</th><th>Gene</th><th>Percentage of coding region covered</th></tr>
        </thead>
        <tbody id="tbody">{table_rows}</tbody>
      </table>
    </div>
  </div>
  <script>
    const searchInput = document.getElementById('searchInput');
    const tbody = document.getElementById('tbody');
    let rows = Array.from(tbody.querySelectorAll('tr'));

    function applyFilters() {{
      const q = searchInput.value.toLowerCase();
      let n = 0;
      rows.forEach(r => {{
        const matchQ = !q || Array.from(r.cells).some(c => c.textContent.toLowerCase().includes(q));
        r.style.display = matchQ ? '' : 'none';
        if (matchQ) {{ n++; r.cells[0].textContent = n; }}
      }});
    }}

    searchInput.addEventListener('keyup', applyFilters);

    function sortTable(ascending) {{
      const sortedRows = rows.slice().sort((a, b) => {{
        const aVal = parseFloat(a.cells[2].textContent) || 0;
        const bVal = parseFloat(b.cells[2].textContent) || 0;
        return ascending ? aVal - bVal : bVal - aVal;
      }});
      sortedRows.forEach(row => tbody.appendChild(row));
      rows = Array.from(tbody.querySelectorAll('tr'));
      applyFilters();
    }}

    document.getElementById('downloadBtn').addEventListener('click', function() {{
      const visibleRows = rows.filter(r => r.style.display !== 'none');
      let csvContent = "S.No,Gene,Percentage\\n";
      visibleRows.forEach(row => {{
        const cells = Array.from(row.cells).map(c => '"' + c.textContent.trim() + '"');
        csvContent += cells.join(',') + '\\n';
      }});
      const blob = new Blob([csvContent], {{ type: 'text/csv;charset=utf-8;' }});
      const link = document.createElement('a');
      link.href = URL.createObjectURL(blob);
      link.download = 'gene_coverage_{safe_patient_js}.csv';
      link.click();
    }});
  </script>
</body>
</html>"""

def filter_panel_genes(coverage_df, panel_genes):
    """Apply the panel-gene filter to one coverage df.
    Returns a grouped df (Gene_ID, Perc_1x) sorted by Gene_ID, or None if no coverage column."""
    df = coverage_df.copy()
    pg = [g.strip().upper() for g in panel_genes]
    if 'Gene_Name' not in df.columns:
        df['Gene_Name'] = df.get('Gene_ID', df.iloc[:, 0]).astype(str)
    if 'Ref Name' not in df.columns:
        df['Ref Name'] = df['Gene_Name']
    df['Gene_Name'] = df['Gene_Name'].astype(str).str.strip().str.upper()
    df['Ref Name']  = df['Ref Name'].astype(str).str.strip().str.upper()
    cov_col = next((c for c in ['% 1x', '%1x'] if c in df.columns), None)
    if not cov_col:
        return None
    fdf = df[df['Gene_Name'].isin(pg) | df['Ref Name'].isin(pg)].copy()
    fdf['Gene_ID'] = fdf['Gene_Name'].fillna(fdf['Ref Name']).str.strip().str.upper()
    fdf = fdf[~fdf['Gene_ID'].fillna('').str.startswith("Intron:")]
    fdf = fdf[~fdf['Gene_ID'].fillna('').str.startswith("ENSG")]
    fdf['Perc_1x'] = pd.to_numeric(fdf[cov_col], errors='coerce')
    g = fdf.groupby('Gene_ID', as_index=False)['Perc_1x'].mean()
    g['Perc_1x'] = g['Perc_1x'].round(2)
    return g.sort_values('Gene_ID')

def build_word_bytes(df_grouped, mito_data=None):
    """Build a Word report (.docx) from a panel-filtered grouped df, optionally appending mito genes."""
    if mito_data is not None:
        df_panel = df_grouped.rename(columns={'Perc_1x': '% 1x'}).sort_values('Gene_ID')
        mito_sorted = mito_data.sort_values('Gene_ID')
        df_final = pd.concat([df_panel, mito_sorted], ignore_index=True)
        df_final = df_final.drop_duplicates(subset='Gene_ID', keep='first').reset_index(drop=True)
    else:
        df_final = df_grouped.rename(columns={'Perc_1x': '% 1x'})
    doc = create_word_document_with_mito(df_final)
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.getvalue()

def build_all_genes_html(coverage_df, mito_data, patient_name):
    """Build the interactive HTML report (ALL coverage genes, mito appended) for one sample."""
    cov = coverage_df.copy()
    gcol = next((c for c in ['Gene_ID', 'Gene_Name', 'Gene'] if c in cov.columns), cov.columns[0])
    cov[gcol] = cov[gcol].astype(str).str.strip().str.upper()
    cov = cov.sort_values(gcol)
    if mito_data is not None:
        mito_s = mito_data.sort_values('Gene_ID').rename(columns={'Gene_ID': gcol})
        cov = pd.concat([cov, mito_s], ignore_index=True).drop_duplicates(subset=gcol, keep='first')
    return generate_html_report(cov, patient_name)

# ──────────────────────────────────────────────────────────────────────────
# Main UI
# ──────────────────────────────────────────────────────────────────────────

# Step 1 — centred pill toggle (no columns — CSS flex centres it at full width)
st.markdown('<div class="gca-step1"><span class="gca-lbl">Step 1 — Select Input Format</span></div>',
            unsafe_allow_html=True)
input_type = st.radio(
    "", ["📊 Pre-processed CSV", "📁 Raw Excel File"],
    horizontal=True, label_visibility="collapsed"
)

st.markdown("<div style='margin-bottom:.75rem'></div>", unsafe_allow_html=True)

# Step 2 — data upload  |  panel genes
col1, col2 = st.columns(2, gap="large")

with col1:
    st.markdown(
        '<div class="gca-card-title">Coverage Data '
        '<span class="gca-badge-ready">Ready to upload</span></div>',
        unsafe_allow_html=True
    )

    if "Raw Excel" in input_type:
        raw_excels = st.file_uploader(
            "Choose Excel file(s)", type=['xlsx', 'xls'],
            accept_multiple_files=True, key='raw_excel'
        )

        if raw_excels:
            # Reprocess only when the uploaded set changes (cache keeps it cheap)
            sig = tuple((f.name, f.size) for f in raw_excels)
            if st.session_state.raw_batch_sig != sig:
                batch, errors = {}, []
                progress = st.progress(0, "Processing…")
                n = len(raw_excels)
                for i, f in enumerate(raw_excels):
                    try:
                        coverage_df, basename = preprocess_excel_cached(f.read(), f.name)
                        batch[basename] = coverage_df
                    except Exception as e:
                        errors.append(f"{f.name}: {e}")
                    progress.progress(int((i + 1) / n * 100), f"Processing {i+1}/{n}…")
                progress.empty()
                for e in errors:
                    st.error(f"❌ {e}")
                st.session_state.raw_batch = batch
                st.session_state.raw_batch_sig = sig
                st.session_state.pipeline_html_bytes = None
                st.session_state.pipeline_html_name  = None

            batch = st.session_state.raw_batch
            if batch:
                names = list(batch.keys())
                if len(names) == 1:
                    active = names[0]
                else:
                    st.success(f"✅ {len(names)} files preprocessed")
                    selected = st.multiselect(
                        "Samples to include in reports (same panel applied to each)",
                        names, default=names, key='raw_selected'
                    )
                    if not selected:
                        selected = names
                    active = st.selectbox(
                        "Preview sample on screen", selected, key='raw_preview'
                    )

                # Switch active sample → reset any stale generated HTML
                if st.session_state.file_basename != active:
                    st.session_state.pipeline_html_bytes = None
                    st.session_state.pipeline_html_name  = None
                st.session_state.coverage_data = batch[active]
                st.session_state.file_basename = active

                st.success(f"✅ {len(batch[active])} records loaded — **{active}**")

                if len(names) > 1:
                    import zipfile
                    zbuf = io.BytesIO()
                    with zipfile.ZipFile(zbuf, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for nm, d in batch.items():
                            zf.writestr(f"{nm}_coverage.csv", d.to_csv(index=False))
                    zbuf.seek(0)
                    st.download_button(
                        f"⬇️ Download all {len(names)} CSVs (ZIP)",
                        data=zbuf.getvalue(),
                        file_name="preprocessed_coverage_csvs.zip",
                        mime="application/zip",
                        use_container_width=True
                    )

                st.info(
                    "💡 **No re-upload needed.** Load panel genes in Step 2B below — "
                    "results appear automatically for the active sample."
                )
                st.download_button(
                    f"⬇️ Download active CSV — {active} (optional)",
                    data=batch[active].to_csv(index=False),
                    file_name=f"{active}_coverage.csv",
                    mime="text/csv",
                    use_container_width=True,
                    key="raw_active_dl"
                )

        with st.expander("➕ Add Mitochondrial Data (optional)"):
            mito_file_raw = st.file_uploader(
                "Mitochondrial Excel (header in row 2)",
                type=['xlsx', 'xls'], key='mito_file_raw'
            )
            if mito_file_raw:
                try:
                    mito_df = pd.read_excel(mito_file_raw, header=1)
                    mito_df.columns = mito_df.columns.str.strip()
                    mito_cols = [c for c in mito_df.columns if '1x' in str(c).lower()]
                    if not mito_cols:
                        st.error("❌ No coverage column in mitochondrial file")
                        st.session_state.mito_data = None
                    elif 'Name' not in mito_df.columns:
                        st.error("❌ No 'Name' column in mitochondrial file")
                        st.session_state.mito_data = None
                    else:
                        mito_df = mito_df.rename(columns={mito_cols[0]: '% 1x'})
                        mito_df['Gene_ID'] = mito_df['Name'].astype(str).str.split('/').str[0].str.strip()
                        mito_df = mito_df[['Gene_ID', '% 1x']].copy()
                        mito_df['% 1x'] = pd.to_numeric(mito_df['% 1x'], errors='coerce')
                        mito_df = mito_df.dropna(subset=['% 1x'])
                        mito_df['% 1x'] = mito_df['% 1x'].round(2)
                        mito_df = mito_df[mito_df['Gene_ID'].str.strip() != '']
                        st.session_state.mito_data = mito_df
                        st.success(f"✅ {len(mito_df)} mitochondrial genes loaded")
                except Exception as e:
                    st.error(f"❌ {e}")
                    st.session_state.mito_data = None
            else:
                st.session_state.mito_data = None
    else:
        coverage_file = st.file_uploader("Choose Coverage CSV", type=['csv'], key='coverage')
        if coverage_file:
            st.session_state.file_basename = os.path.splitext(coverage_file.name)[0]
            df = pd.read_csv(coverage_file)
            df.columns = df.columns.str.strip()
            coverage_col = next((c for c in ['% 1x', '%1x'] if c in df.columns), None)
            if coverage_col:
                st.session_state.coverage_data = df
                st.session_state.pipeline_html_bytes = None
                st.session_state.pipeline_html_name  = None
                st.success(f"✅ {len(df)} records loaded")
            else:
                st.error("❌ No coverage column found ('% 1x' or '%1x')")

        with st.expander("➕ Add Mitochondrial Data (optional)"):
            mito_file = st.file_uploader(
                "Mitochondrial Excel (header in row 2)",
                type=['xlsx', 'xls'], key='mito_file'
            )
            if mito_file:
                try:
                    mito_df = pd.read_excel(mito_file, header=1)
                    mito_df.columns = mito_df.columns.str.strip()
                    mito_cols = [c for c in mito_df.columns if '1x' in str(c).lower()]
                    if not mito_cols:
                        st.error("❌ No coverage column in mitochondrial file")
                        st.session_state.mito_data = None
                    elif 'Name' not in mito_df.columns:
                        st.error("❌ No 'Name' column in mitochondrial file")
                        st.session_state.mito_data = None
                    else:
                        mito_df = mito_df.rename(columns={mito_cols[0]: '% 1x'})
                        mito_df['Gene_ID'] = mito_df['Name'].astype(str).str.split('/').str[0].str.strip()
                        mito_df = mito_df[['Gene_ID', '% 1x']].copy()
                        mito_df['% 1x'] = pd.to_numeric(mito_df['% 1x'], errors='coerce')
                        before = len(mito_df)
                        mito_df = mito_df.dropna(subset=['% 1x'])
                        if before > len(mito_df):
                            st.warning(f"⚠️ Removed {before - len(mito_df)} rows with invalid coverage")
                        mito_df['% 1x'] = mito_df['% 1x'].round(2)
                        mito_df = mito_df[mito_df['Gene_ID'].str.strip() != '']
                        st.session_state.mito_data = mito_df
                        st.success(f"✅ {len(mito_df)} mitochondrial genes loaded")
                except Exception as e:
                    st.error(f"❌ {e}")
                    import traceback; st.code(traceback.format_exc())
                    st.session_state.mito_data = None
            else:
                st.session_state.mito_data = None

with col2:
    st.markdown('<div class="gca-card-title">Target Panel Genes</div>', unsafe_allow_html=True)
    tab1, tab2 = st.tabs(["📁 Upload Excel", "📝 Paste Genes"])

    with tab1:
        panel_file = st.file_uploader("Choose Excel", type=['xlsx', 'xls'], key='panel')
        if panel_file:
            panel_df = pd.read_excel(panel_file)
            if 'GENE' in panel_df.columns:
                genes = panel_df['GENE'].dropna().unique().tolist()
                st.session_state.panel_genes = genes
                st.success(f"✅ {len(genes)} genes loaded")
            else:
                st.error("❌ No 'GENE' column found")

    with tab2:
        gene_text = st.text_area(
            "Gene symbols",
            height=150,
            placeholder="BRCA1 BRCA2 TP53\nEGFR, KRAS\nTP53",
            label_visibility="collapsed"
        )
        c_left, c_right = st.columns([2, 1])
        with c_right:
            if st.button("Load Genes", use_container_width=True):
                genes, n_dup = parse_gene_list(gene_text)
                if genes:
                    st.session_state.panel_genes = genes
                    msg = f"✅ {len(genes)} genes"
                    if n_dup:
                        msg += f" ({n_dup} duplicates removed)"
                    st.success(msg)
                else:
                    st.warning("⚠️ No genes found")
        with c_left:
            if st.session_state.panel_genes:
                st.markdown(
                    f'<span style="font-size:12px;color:var(--cs)">'
                    f'● {len(st.session_state.panel_genes)} genes loaded</span>',
                    unsafe_allow_html=True
                )

# ── Results ────────────────────────────────────────────────────────────────
if st.session_state.coverage_data is not None and st.session_state.panel_genes:
    panel_genes = st.session_state.panel_genes

    # Samples to report on: in Raw-Excel multi mode = the multiselect subset; else the single active sample
    if "Raw Excel" in input_type and len(st.session_state.raw_batch) > 1:
        sel = st.session_state.get('raw_selected') or list(st.session_state.raw_batch.keys())
        report_samples = {nm: st.session_state.raw_batch[nm] for nm in sel if nm in st.session_state.raw_batch}
    else:
        report_samples = {}
    if not report_samples:
        report_samples = {st.session_state.file_basename: st.session_state.coverage_data}
    multi = len(report_samples) > 1

    # On-screen preview uses the active sample (coverage_data)
    df_grouped = filter_panel_genes(st.session_state.coverage_data, panel_genes)
    if df_grouped is None:
        st.error("Coverage column not found")
        st.stop()
    st.session_state.filtered_data = df_grouped

    st.divider()

    total = len(df_grouped)
    low   = len(df_grouped[df_grouped['Perc_1x'] < 90])
    avg   = df_grouped['Perc_1x'].mean()

    st.markdown(f"""
<div class="gca-rh">
  <div>
    <div class="gca-rh-ttl">Analysis Results</div>
    <div class="gca-rh-sub">Sample: {html_lib.escape(st.session_state.file_basename)}</div>
  </div>
</div>
<div class="gca-metrics">
  <div class="gca-metric">
    <div class="gca-metric-lbl">Total Genes Analyzed</div>
    <div class="gca-metric-val">{total}</div>
  </div>
  <div class="gca-metric err">
    <div class="gca-metric-lbl">Low Coverage (&lt; 90%)</div>
    <div class="gca-metric-val">{low}</div>
  </div>
  <div class="gca-metric">
    <div class="gca-metric-lbl">Average Coverage</div>
    <div class="gca-metric-val">{avg:.1f}%</div>
  </div>
</div>
""", unsafe_allow_html=True)

    if multi:
        st.caption(f"📦 Reports will be generated for **{len(report_samples)} selected samples** "
                   "(same panel applied to each), bundled as ZIP files.")

    has_mito = st.session_state.mito_data is not None
    mito_for_report = st.session_state.mito_data if has_mito else None
    btn_col1, btn_col2 = st.columns(2, gap="medium")

    with btn_col1:
        if not multi:
            if st.button("📄 Generate Word Report", type="primary", use_container_width=True):
                with st.spinner("Creating…"):
                    try:
                        word_bytes = build_word_bytes(df_grouped, mito_for_report)
                        st.session_state.doc_bytes     = word_bytes
                        st.session_state.word_filename = f"{st.session_state.file_basename}_report.docx"
                        st.success("✅ Ready to download!")
                    except Exception as e:
                        st.error(f"❌ {e}")
                        import traceback; st.code(traceback.format_exc())

            if st.session_state.get('doc_bytes'):
                st.download_button(
                    f"⬇️ {st.session_state.word_filename}",
                    data=st.session_state.doc_bytes,
                    file_name=st.session_state.word_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        else:
            if st.button(f"📄 Generate Word Reports ({len(report_samples)})",
                         type="primary", use_container_width=True, key="batch_word_btn"):
                with st.spinner("Creating Word reports…"):
                    try:
                        import zipfile
                        zbuf = io.BytesIO()
                        with zipfile.ZipFile(zbuf, 'w', zipfile.ZIP_DEFLATED) as zf:
                            for nm, cov in report_samples.items():
                                g = filter_panel_genes(cov, panel_genes)
                                if g is None or g.empty:
                                    continue
                                zf.writestr(f"{nm}_report.docx", build_word_bytes(g, mito_for_report))
                        zbuf.seek(0)
                        st.session_state.batch_word_zip = zbuf.getvalue()
                        st.success("✅ Ready to download!")
                    except Exception as e:
                        st.error(f"❌ {e}")
                        import traceback; st.code(traceback.format_exc())

            if st.session_state.get('batch_word_zip'):
                st.download_button(
                    "⬇️ Download Word Reports (ZIP)",
                    data=st.session_state.batch_word_zip,
                    file_name="word_reports.zip",
                    mime="application/zip",
                    use_container_width=True,
                    key="batch_word_dl",
                )

    with btn_col2:
        st.markdown('<div class="gca-sec">', unsafe_allow_html=True)
        if not multi:
            st.download_button(
                "📊 Download Filtered CSV",
                data=df_grouped.to_csv(index=False),
                file_name=f"{st.session_state.file_basename}_filtered.csv",
                mime="text/csv",
                use_container_width=True
            )
        else:
            import zipfile
            zbuf = io.BytesIO()
            with zipfile.ZipFile(zbuf, 'w', zipfile.ZIP_DEFLATED) as zf:
                for nm, cov in report_samples.items():
                    g = filter_panel_genes(cov, panel_genes)
                    if g is None:
                        continue
                    zf.writestr(f"{nm}_filtered.csv", g.to_csv(index=False))
            zbuf.seek(0)
            st.download_button(
                f"📊 Download Filtered CSVs — ZIP ({len(report_samples)})",
                data=zbuf.getvalue(),
                file_name="filtered_csvs.zip",
                mime="application/zip",
                use_container_width=True
            )
        st.markdown('</div>', unsafe_allow_html=True)

    with st.expander("🔬 View Gene Level Details"):
        show_low_only = st.toggle("Show only genes with < 90% coverage", value=False)
        df_display = df_grouped[df_grouped['Perc_1x'] < 90] if show_low_only else df_grouped
        if show_low_only:
            st.caption(f"{len(df_display)} gene(s) with coverage below 90%")
        else:
            st.caption("Red badges = coverage below 90%")
        badges = '<div class="gca-gene-grid">'
        for _, row in df_display.iterrows():
            pct  = row['Perc_1x']
            gene = html_lib.escape(str(row['Gene_ID']))
            cls  = 'low' if pct < 90 else 'ok'
            badges += (
                f'<span class="gca-gene {cls}" title="{pct}%">'
                f'{gene}<span class="pct">({pct}%)</span></span>'
            )
        badges += '</div>'
        st.markdown(badges, unsafe_allow_html=True)

# ── Pipeline shortcut: direct HTML export (needs only coverage data) ────────
if st.session_state.coverage_data is not None:
    st.divider()
    direct_html_on = st.checkbox(
        "🔗 **Direct HTML export** — generate the interactive HTML report from current results "
        "(skip manual CSV download → re-upload cycle)",
        key="pipeline_direct_html",
        value=False,
    )

    if direct_html_on:
        has_mito_p = st.session_state.mito_data is not None
        mito_p = st.session_state.mito_data if has_mito_p else None

        # Samples for HTML: Raw-Excel multi mode = the multiselect subset; else the single active sample
        if "Raw Excel" in input_type and len(st.session_state.raw_batch) > 1:
            sel_h = st.session_state.get('raw_selected') or list(st.session_state.raw_batch.keys())
            html_samples = {nm: st.session_state.raw_batch[nm] for nm in sel_h if nm in st.session_state.raw_batch}
        else:
            html_samples = {}
        if not html_samples:
            html_samples = {st.session_state.file_basename: st.session_state.coverage_data}
        multi_h = len(html_samples) > 1

        col_gen, col_info = st.columns([2, 3])
        with col_info:
            mito_note = " + mitochondrial genes at end" if has_mito_p else ""
            if multi_h:
                st.caption(f"Will export **all genes** for **{len(html_samples)} selected samples**{mito_note} (ZIP).")
            else:
                all_gene_count = len(st.session_state.coverage_data)
                st.caption(f"Will export **all {all_gene_count} genes** from coverage data{mito_note}.")

        with col_gen:
            if st.button("🌐 Build HTML Report", use_container_width=True, type="primary",
                         key="pipeline_build_btn"):
                with st.spinner("Building report…"):
                    if multi_h:
                        import zipfile
                        zbuf = io.BytesIO()
                        with zipfile.ZipFile(zbuf, 'w', zipfile.ZIP_DEFLATED) as zf:
                            for nm, cov in html_samples.items():
                                html_out = build_all_genes_html(cov, mito_p, nm)
                                safe_nm = nm.replace(' ', '_').replace('/', '_')
                                zf.writestr(f"coverage_{safe_nm}.html", html_out)
                        zbuf.seek(0)
                        st.session_state.pipeline_html_bytes = zbuf.getvalue()
                        st.session_state.pipeline_html_name  = None  # signals ZIP
                    else:
                        nm = st.session_state.file_basename
                        html_out = build_all_genes_html(st.session_state.coverage_data, mito_p, nm)
                        st.session_state.pipeline_html_bytes = html_out.encode('utf-8')
                        st.session_state.pipeline_html_name  = nm.replace(' ', '_').replace('/', '_')

        if st.session_state.pipeline_html_bytes:
            st.success("✅ HTML ready — download below.")
            if st.session_state.pipeline_html_name is None:
                st.download_button(
                    f"⬇️ Download HTML Reports — ZIP ({len(html_samples)})",
                    data=st.session_state.pipeline_html_bytes,
                    file_name="html_reports.zip",
                    mime="application/zip",
                    use_container_width=True,
                    key="pipeline_html_dl",
                )
            else:
                st.download_button(
                    f"⬇️ Download coverage_{st.session_state.pipeline_html_name}.html",
                    data=st.session_state.pipeline_html_bytes,
                    file_name=f"coverage_{st.session_state.pipeline_html_name}.html",
                    mime="text/html",
                    use_container_width=True,
                    key="pipeline_html_dl",
                )

st.divider()

# ── Batch HTML Generator ───────────────────────────────────────────────────
with st.expander("📦 Batch HTML Report Generator"):
    st.markdown(
        "Upload multiple CSV files to get interactive HTML reports — "
        "each with search, sorting, and download features."
    )

    batch_files = st.file_uploader(
        "Upload CSV files",
        type=['csv'],
        accept_multiple_files=True,
        key='batch_csvs'
    )

    if batch_files:
        st.info(f"✅ {len(batch_files)} file(s) loaded")

        if st.button("🎨 Generate HTML Reports", type="primary", use_container_width=True):
            pb = st.progress(0, "Generating…")
            try:
                from datetime import datetime
                n = len(batch_files)

                if n == 1:
                    csv_file    = batch_files[0]
                    patient_name = os.path.splitext(csv_file.name)[0]
                    df           = pd.read_csv(csv_file)
                    html_content = generate_html_report(df, patient_name)
                    pb.progress(1.0, "Done!"); pb.empty()
                    st.success("✅ Report ready!")
                    safe_name = patient_name.replace(' ', '_').replace('/', '_')
                    st.download_button(
                        f"⬇️ Download {safe_name}.html",
                        data=html_content.encode('utf-8'),
                        file_name=f"coverage_{safe_name}.html",
                        mime="text/html",
                        use_container_width=True, type="primary"
                    )
                else:
                    import zipfile
                    zip_buf = io.BytesIO()
                    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for idx, csv_file in enumerate(batch_files):
                            patient_name = os.path.splitext(csv_file.name)[0]
                            df           = pd.read_csv(csv_file)
                            html_content = generate_html_report(df, patient_name)
                            safe_name    = patient_name.replace(' ', '_').replace('/', '_')
                            zf.writestr(f"coverage_{safe_name}.html", html_content)
                            pb.progress((idx + 1) / n, f"Creating {idx+1}/{n}…")
                    zip_buf.seek(0); pb.empty()
                    st.success(f"✅ {n} reports generated!")
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    st.download_button(
                        "⬇️ Download All Reports (ZIP)",
                        data=zip_buf.getvalue(),
                        file_name=f"gene_coverage_reports_{ts}.zip",
                        mime="application/zip",
                        use_container_width=True, type="primary"
                    )
            except Exception as e:
                pb.empty()
                st.error(f"❌ {e}")
                import traceback; st.code(traceback.format_exc())

# ── Footer ─────────────────────────────────────────────────────────────────
st.markdown("""
<div class="gca-footer">
  <span class="ms">lock</span>
  All processing happens on the server. Data is not stored permanently.
</div>
""", unsafe_allow_html=True)
